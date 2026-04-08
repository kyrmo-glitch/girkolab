from appJar import gui
from docx import Document
from figure_pkg import *

last_rect_result = ""
last_tri_result = ""
last_trap_result = ""

def calculate_rectangle(btn):
    global last_rect_result
    
    a = float(app.getEntry("rect_a"))
    b = float(app.getEntry("rect_b"))
    
    s = rect_area(a, b)
    ri = rect_r_in(a, b)
    ro = rect_r_out(a, b)
    
    text = f"Площадь: {round(s,2)}\n"
    if ri > 0:
        text += f"Радиус вписанной: {round(ri,2)}\n"
    else:
        text += "Вписанной окружности нет\n"
    text += f"Радиус описанной: {round(ro,2)}"
    
    app.setLabel("rect_result", text)
    last_rect_result = text

def calculate_triangle(btn):
    global last_tri_result
    
    a = float(app.getEntry("tri_a"))
    b = float(app.getEntry("tri_b"))
    c = float(app.getEntry("tri_c"))
    
    s = tri_area(a, b, c)
    ri = tri_r_in(a, b, c)
    ro = tri_r_out(a, b, c)
    
    text = f"Площадь: {round(s,2)}\n"
    text += f"Радиус вписанной: {round(ri,2)}\n"
    text += f"Радиус описанной: {round(ro,2)}"
    
    app.setLabel("tri_result", text)
    last_tri_result = text

def calculate_trapezoid(btn):
    global last_trap_result
    
    a = float(app.getEntry("trap_a"))
    b = float(app.getEntry("trap_b"))
    h = float(app.getEntry("trap_h"))
    
    s = trap_area(a, b, h)
    
    text = f"Площадь: {round(s,2)}\n"
    text += "Радиусы не рассчитываются"
    
    app.setLabel("trap_result", text)
    last_trap_result = text

def save_all(btn):
    """Сохранение в Word"""
    doc = Document()
    doc.add_heading("Расчёт геометрических фигур", 0)
    
    if last_rect_result:
        doc.add_paragraph(last_rect_result)
    if last_tri_result:
        doc.add_paragraph(last_tri_result)
    if last_trap_result:
        doc.add_paragraph(last_trap_result)
    
    file_path = app.saveBox(title="Сохранить файл",
                            fileName="результаты.docx",
                            fileExt=[("Word документ", "*.docx"), ("Все файлы", "*.*")])
    
    if file_path:
        if not file_path.endswith('.docx'):
            file_path += '.docx'
        doc.save(file_path)
        app.infoBox("Успех", "Файл сохранён!")

# Создание GUI
app = gui("Калькулятор геометрических фигур", "1100x900")
app.setFont(12)

# Прямоугольник
app.addLabel("rect_section", "ПРЯМОУГОЛЬНИК", 1, 0)

app.addLabel("rect_a_label", "Сторона a:", 2, 0)
app.addEntry("rect_a", 2, 1,1)

app.addLabel("rect_b_label", "Сторона b:", 3, 0)
app.addEntry("rect_b", 3, 1,1)

app.addButton("Рассчитать прямоугольник", calculate_rectangle, 4, 0, 2)

app.addLabel("rect_result_label", "Результат:", 5, 0)
app.addLabel("rect_result", "Результат", 5, 1)

# Треугольник
app.addHorizontalSeparator(6, 0, 2)
app.setLabelFont("tri_section", size=14)
app.addLabel("tri_section", "ТРЕУГОЛЬНИК", 7, 0, 2)

app.addLabel("tri_a_label", "Сторона a:", 8, 0)
app.addEntry("tri_a", 8, 1)

app.addLabel("tri_b_label", "Сторона b:", 9, 0)
app.addEntry("tri_b", 9, 1)

app.addLabel("tri_c_label", "Сторона c:", 10, 0)
app.addEntry("tri_c", 10, 1)

app.addButton("Рассчитать треугольник", calculate_triangle, 11, 0, 2)

app.addLabel("tri_result_label", "Результат:", 12, 0)
app.addLabel("tri_result", "Результат", 12, 1)

# Трапеция
app.addHorizontalSeparator(13, 0, 2)
app.addLabel("trap_section", "ТРАПЕЦИЯ", 14, 0, 2)

app.addLabel("trap_a_label", "Основание a:", 15, 0)
app.addEntry("trap_a", 15, 1)

app.addLabel("trap_b_label", "Основание b:", 16, 0)
app.addEntry("trap_b", 16, 1)

app.addLabel("trap_h_label", "Высота h:", 17, 0)
app.addEntry("trap_h", 17, 1)

app.addButton("Рассчитать трапецию", calculate_trapezoid, 18, 0, 2)

app.addLabel("trap_result_label", "Результат:", 19, 0)
app.addLabel("trap_result", "Результат", 19, 1)

# Кнопка сохранения
app.addHorizontalSeparator(20, 0, 2)
app.addButton("Сохранить в Word", save_all, 21, 0,2)

# Запуск приложения
app.go()