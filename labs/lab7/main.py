import tkinter as tk
from tkinter import messagebox, filedialog
from abc import ABC, abstractmethod
from docx import Document
from figure_pkg import *


class Figure(ABC):
    """Абстрактный базовый класс"""
    
    def __init__(self, name):
        self._name = name
        self._area = 0
        self._result = ""
    
    @property
    def area(self):
        return self._area
    
    @area.setter
    def area(self, value):
        self._area = value
    
    @abstractmethod
    def calculate(self):
        pass
    
    def __str__(self):
        return f"{self._name}: {self._result}"
    
    def __len__(self):
        return len(self._result)


class Rectangle(Figure):
    def __init__(self):
        super().__init__("Прямоугольник")
        self.a = 0
        self.b = 0
    
    def calculate(self):
        self.area = rect_area(self.a, self.b)
        ri = rect_r_in(self.a, self.b)
        ro = rect_r_out(self.a, self.b)
        
        self._result = f"Площадь: {round(self.area, 2)}\n"
        if ri > 0:
            self._result += f"Радиус вписанной: {round(ri, 2)}\n"
        else:
            self._result += "Вписанной окружности нет\n"
        self._result += f"Радиус описанной: {round(ro, 2)}"
        return self._result
    
    def __repr__(self):
        return f"Rectangle({self.a}, {self.b})"
    
    def __eq__(self, other):
        return self.area == other.area if isinstance(other, Figure) else False


class Triangle(Figure):
    def __init__(self):
        super().__init__("Треугольник")
        self.a = 0
        self.b = 0
        self.c = 0
    
    def calculate(self):
        self.area = tri_area(self.a, self.b, self.c)
        ri = tri_r_in(self.a, self.b, self.c)
        ro = tri_r_out(self.a, self.b, self.c)
        
        self._result = f"Площадь: {round(self.area, 2)}\n"
        self._result += f"Радиус вписанной: {round(ri, 2)}\n"
        self._result += f"Радиус описанной: {round(ro, 2)}"
        return self._result
    
    def __repr__(self):
        return f"Triangle({self.a}, {self.b}, {self.c})"
    
    def __eq__(self, other):
        return self.area == other.area if isinstance(other, Figure) else False


class Trapezoid(Figure):
    def __init__(self):
        super().__init__("Трапеция")
        self.a = 0
        self.b = 0
        self.h = 0
    
    def calculate(self):
        self.area = trap_area(self.a, self.b, self.h)
        self._result = f"Площадь: {round(self.area, 2)}\n"
        self._result += "Радиусы не рассчитываются"
        return self._result
    
    def __repr__(self):
        return f"Trapezoid({self.a}, {self.b}, {self.h})"
    
    def __eq__(self, other):
        return self.area == other.area if isinstance(other, Figure) else False


class App:
    def __init__(self):
        self.root = tk.Tk()
        self.root.title("Калькулятор фигур")
        self.root.geometry("500x750")
        
        self.rect = Rectangle()
        self.tri = Triangle()
        self.trap = Trapezoid()
        
        self.setup_ui()
    
    def setup_ui(self):
        # ===== ПРЯМОУГОЛЬНИК =====
        tk.Label(self.root, text="ПРЯМОУГОЛЬНИК", font=("Arial", 14, "bold")).pack(pady=10)
        
        tk.Label(self.root, text="Сторона a:").pack()
        self.rect_a = tk.Entry(self.root)
        self.rect_a.pack()
        
        tk.Label(self.root, text="Сторона b:").pack()
        self.rect_b = tk.Entry(self.root)
        self.rect_b.pack()
        
        tk.Button(self.root, text="Рассчитать", command=self.calc_rectangle).pack(pady=5)
        self.rect_result = tk.Label(self.root, text="Результат:", bg="lightgray", relief=tk.SUNKEN)
        self.rect_result.pack(fill=tk.X, padx=20, pady=5)
        
        # ===== ТРЕУГОЛЬНИК =====
        tk.Label(self.root, text="ТРЕУГОЛЬНИК", font=("Arial", 14, "bold")).pack(pady=10)
        
        tk.Label(self.root, text="Сторона a:").pack()
        self.tri_a = tk.Entry(self.root)
        self.tri_a.pack()
        
        tk.Label(self.root, text="Сторона b:").pack()
        self.tri_b = tk.Entry(self.root)
        self.tri_b.pack()
        
        tk.Label(self.root, text="Сторона c:").pack()
        self.tri_c = tk.Entry(self.root)
        self.tri_c.pack()
        
        tk.Button(self.root, text="Рассчитать", command=self.calc_triangle).pack(pady=5)
        self.tri_result = tk.Label(self.root, text="Результат:", bg="lightgray", relief=tk.SUNKEN)
        self.tri_result.pack(fill=tk.X, padx=20, pady=5)
        
        # ===== ТРАПЕЦИЯ =====
        tk.Label(self.root, text="ТРАПЕЦИЯ", font=("Arial", 14, "bold")).pack(pady=10)
        
        tk.Label(self.root, text="Основание a:").pack()
        self.trap_a = tk.Entry(self.root)
        self.trap_a.pack()
        
        tk.Label(self.root, text="Основание b:").pack()
        self.trap_b = tk.Entry(self.root)
        self.trap_b.pack()
        
        tk.Label(self.root, text="Высота h:").pack()
        self.trap_h = tk.Entry(self.root)
        self.trap_h.pack()
        
        tk.Button(self.root, text="Рассчитать", command=self.calc_trapezoid).pack(pady=5)
        self.trap_result = tk.Label(self.root, text="Результат:", bg="lightgray", relief=tk.SUNKEN)
        self.trap_result.pack(fill=tk.X, padx=20, pady=5)
        
        # ===== СОХРАНЕНИЕ =====
        tk.Button(self.root, text="Сохранить в Word", command=self.save_to_word, bg="lightblue", font=("Arial", 12)).pack(pady=20)
    
    def calc_rectangle(self):
        self.rect.a = float(self.rect_a.get())
        self.rect.b = float(self.rect_b.get())
        result = self.rect.calculate()
        self.rect_result.config(text=f"Результат:\n{result}")
    
    def calc_triangle(self):
        self.tri.a = float(self.tri_a.get())
        self.tri.b = float(self.tri_b.get())
        self.tri.c = float(self.tri_c.get())
        result = self.tri.calculate()
        self.tri_result.config(text=f"Результат:\n{result}")
    
    def calc_trapezoid(self):
        self.trap.a = float(self.trap_a.get())
        self.trap.b = float(self.trap_b.get())
        self.trap.h = float(self.trap_h.get())
        result = self.trap.calculate()
        self.trap_result.config(text=f"Результат:\n{result}")
    
    def save_to_word(self):
        doc = Document()
        doc.add_heading("Результаты расчётов", 0)
        
        for fig in [self.rect, self.tri, self.trap]:
            if fig._result:
                doc.add_paragraph(fig._result)
                doc.add_paragraph(f"str: {str(fig)}")
                doc.add_paragraph(f"len: {len(fig)} символов")
                doc.add_paragraph("-" * 40)
        
        if self.rect.area > 0 and self.tri.area > 0:
            doc.add_paragraph(f"Площади равны? {self.rect == self.tri}")
        
        path = filedialog.asksaveasfilename(defaultextension=".docx", filetypes=[("Word", "*.docx")], initialfile="результаты.docx")
        if path:
            doc.save(path)
            messagebox.showinfo("Успех", "Сохранено!")
    
    def run(self):
        self.root.mainloop()


if __name__ == "__main__":
    app = App()
    app.run()