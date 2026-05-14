#docker-compose up -d
#SELECT * FROM results;
#SELECT * FROM results ORDER BY id DESC LIMIT 5;
import tkinter as tk
from tkinter import messagebox, filedialog
from abc import ABC, abstractmethod
from docx import Document
from figure_pkg import *
import psycopg2
from datetime import *

try:
    conn = psycopg2.connect(
        host="localhost",
        database="calculator_db",
        user="user",
        password="password"
    )
    cursor = conn.cursor()
    
    cursor.execute("""
        CREATE TABLE IF NOT EXISTS results (
            id SERIAL PRIMARY KEY,
            figure_type TEXT,
            params TEXT,
            area REAL,
            full_result TEXT,
            date TIMESTAMP DEFAULT CURRENT_TIMESTAMP
        )
    """)
    conn.commit()
    print("База данных подключена")
except Exception as e:
    print(f"Ошибка БД: {e}")
    conn = None
    cursor = None



class Figure(ABC):
    def __init__(self, name):
        self._name = name
        self._area = 0
        self._result = ""
    
    @property
    def area(self):
        return self._area
    
    @area.setter
    def area(self, value):
        if value < 0:                    
            raise ValueError("Площадь не может быть отрицательной!")
        self._area = value
    
    @abstractmethod
    def calculate(self):
        pass
    
    def __str__(self):
        result = f"{self._result}"
        print(f"[__str__] {self._name}")
        print(result)
        return result
    
    def __len__(self):
        length = len(self._result)
        print(f"[__len__] {self._name}")
        print(f"Длина = {length}")
        return length


class Rectangle(Figure):
    def __init__(self):
        super().__init__("Прямоугольник")
        self.a = 0
        self.b = 0
    
    def calculate(self):
        self.area = rect_area(self.a, self.b)
        ri = rect_r_in(self.a, self.b)
        ro = rect_r_out(self.a, self.b)
        
        self._result = f"Площадь: {round(self.area, 2)}\n"
        if ri > 0:
            self._result += f"Радиус вписанной: {round(ri, 2)}\n"
        else:
            self._result += "Вписанной окружности нет\n"
        self._result += f"Радиус описанной: {round(ro, 2)}"
        if conn:
            cursor.execute(
                "INSERT INTO results (figure_type, params, area, full_result) VALUES (%s, %s, %s, %s)",
                ("Прямоугольник", f"a={self.a}, b={self.b}", self.area, self._result)
            )
            conn.commit()
        return self._result
    
    
    def __repr__(self):
        print(f"[__repr__] Rectangle({self.a}, {self.b})")
        return f"Rectangle({self.a}, {self.b})"
    
    def __add__(self, other):
        if isinstance(other, Figure):
            result = self.area + other.area
            print(f"[__add__] {self.area} + {other.area} = {result}")
            return result
        return None


class Triangle(Figure):
    def __init__(self):
        super().__init__("Треугольник")
        self.a = 0
        self.b = 0
        self.c = 0
    
    def calculate(self):
        self.area = tri_area(self.a, self.b, self.c)
        ri = tri_r_in(self.a, self.b, self.c)
        ro = tri_r_out(self.a, self.b, self.c)
        
        self._result = f"Площадь: {round(self.area, 2)}\n"
        self._result += f"Радиус вписанной: {round(ri, 2)}\n"
        self._result += f"Радиус описанной: {round(ro, 2)}"
        if conn:
            cursor.execute(
                "INSERT INTO results (figure_type, params, area, full_result) VALUES (%s, %s, %s, %s)",
                ("Треугольник", f"a={self.a}, b={self.b}, c={self.c}", self.area, self._result)
            )
            conn.commit()
        return self._result
    
    def __repr__(self):
        print(f"[__repr__] Triangle({self.a}, {self.b}, {self.c})")
        return f"Triangle({self.a}, {self.b}, {self.c})"
    
    def __sub__(self, other):
        if isinstance(other, Figure):
            result = abs(self.area - other.area)
            print(f"[__sub__] |{self.area} - {other.area}| = {result}")
            return result
        return None


class Trapezoid(Figure):
    def __init__(self):
        super().__init__("Трапеция")
        self.a = 0
        self.b = 0
        self.h = 0
    
    def calculate(self):
        """Расчёт параметров трапеции"""
        self.area = trap_area(self.a, self.b, self.h)
        ri = trap_r_in(self.a, self.b, self.h)
        ro = trap_r_out(self.a, self.b, self.h)  
        self._result = f"Площадь: {round(self.area, 2)}\n"
        if ri is not None and ri > 0:
            self._result += f"Радиус вписанной: {round(ri, 2)}\n"
        else:
            self._result += "Вписанной окружности нет\n"
        if ro is not None and ro > 0:
            self._result += f"Радиус описанной: {round(ro, 2)}\n"
        else:
            self._result += "Описанной окружности нет\n"
        if conn:
            cursor.execute(
                "INSERT INTO results (figure_type, params, area, full_result) VALUES (%s, %s, %s, %s)",
                ("Трапеция", f"a={self.a}, b={self.b}, h={self.h}", self.area, self._result)
            )
            conn.commit()
        return self._result
    
    def __repr__(self):
        print(f"[__repr__] Trapezoid({self.a}, {self.b}, {self.h})")
        return f"Trapezoid({self.a}, {self.b}, {self.h})"
    
    def __mul__(self, factor):
        if isinstance(factor, (int, float)):
            result = self.area * factor
            print(f"[__mul__] {self.area} * {factor} = {result}")
            return result
        return None


class App:
    
    def __init__(self):
        self.root = tk.Tk()
        self.root.title("Калькулятор фигур")
        self.root.geometry("1000x1000")
        
        self.rect = Rectangle()
        self.tri = Triangle()
        self.trap = Trapezoid()
        
        self.setup_ui()
    
    def setup_ui(self):
        tk.Label(self.root, text="ПРЯМОУГОЛЬНИК", font=("Arial", 14, "bold")).pack(pady=10)
        
        tk.Label(self.root, text="Сторона a:").pack()
        self.rect_a = tk.Entry(self.root)
        self.rect_a.pack()
        
        tk.Label(self.root, text="Сторона b:").pack()
        self.rect_b = tk.Entry(self.root)
        self.rect_b.pack()
        
        tk.Button(self.root, text="Рассчитать", command=self.calc_rectangle).pack(pady=5)
        self.rect_result = tk.Label(self.root, text="Результат:", bg="lightgray", relief=tk.SUNKEN)
        self.rect_result.pack(fill=tk.X, padx=20, pady=5)
        
        tk.Label(self.root, text="ТРЕУГОЛЬНИК", font=("Arial", 14, "bold")).pack(pady=10)
        
        tk.Label(self.root, text="Сторона a:").pack()
        self.tri_a = tk.Entry(self.root)
        self.tri_a.pack()
        
        tk.Label(self.root, text="Сторона b:").pack()
        self.tri_b = tk.Entry(self.root)
        self.tri_b.pack()
        
        tk.Label(self.root, text="Сторона c:").pack()
        self.tri_c = tk.Entry(self.root)
        self.tri_c.pack()
        
        tk.Button(self.root, text="Рассчитать", command=self.calc_triangle).pack(pady=5)
        self.tri_result = tk.Label(self.root, text="Результат:", bg="lightgray", relief=tk.SUNKEN)
        self.tri_result.pack(fill=tk.X, padx=20, pady=5)
        
        tk.Label(self.root, text="ТРАПЕЦИЯ", font=("Arial", 14, "bold")).pack(pady=10)
        
        tk.Label(self.root, text="Основание a:").pack()
        self.trap_a = tk.Entry(self.root)
        self.trap_a.pack()
        
        tk.Label(self.root, text="Основание b:").pack()
        self.trap_b = tk.Entry(self.root)
        self.trap_b.pack()
        
        tk.Label(self.root, text="Высота h:").pack()
        self.trap_h = tk.Entry(self.root)
        self.trap_h.pack()
        
        tk.Button(self.root, text="Рассчитать", command=self.calc_trapezoid).pack(pady=5)
        self.trap_result = tk.Label(self.root, text="Результат:", bg="lightgray", relief=tk.SUNKEN)
        self.trap_result.pack(fill=tk.X, padx=20, pady=5)
        
        tk.Button(self.root, text="Сохранить в Word", command=self.save_to_word,bg="lightblue", font=("Arial", 12)).pack(pady=20)
        
        tk.Button(self.root, text="Показать историю", command=self.show_history,bg="lightyellow", font=("Arial", 12)).pack(pady=5)
    
    def show_history(self):
        if conn:
            cursor.execute("SELECT id, figure_type, params, area, date FROM results ORDER BY id DESC LIMIT 10")
            results = cursor.fetchall()
        
            if results:
                history = "Последние 10 расчетов:\n\n"
                for row in results:
                    history += f"ID:{row[0]} | {row[1]} | {row[2]} | Площадь: {row[3]} | {row[4]}\n"
                messagebox.showinfo("История расчетов", history)
            else:
                messagebox.showinfo("История", "Нет сохраненных расчетов")
        else:
            messagebox.showerror("Ошибка", "Нет подключения к БД")
    
    def calc_rectangle(self):
        self.rect.a = float(self.rect_a.get())
        self.rect.b = float(self.rect_b.get())
        result = self.rect.calculate()
        self.rect_result.config(text=f"Результат:\n{result}")
        
        str(self.rect)
        len(self.rect)
        repr(self.rect)
        if self.rect.area > 0 and self.tri.area > 0 and self.trap.area > 0:
            print("\n" + "-"*50)
            print("ОПЕРАЦИИ С ФИГУРАМИ")
            print("-"*50)
            self.rect + self.tri
            self.tri - self.trap
            self.trap * 2
            
           
    def calc_triangle(self):
        self.tri.a = float(self.tri_a.get())
        self.tri.b = float(self.tri_b.get())
        self.tri.c = float(self.tri_c.get())
        result = self.tri.calculate()
        self.tri_result.config(text=f"Результат:\n{result}")
            
        str(self.tri)
        len(self.tri)
        repr(self.tri)
        if self.rect.area > 0 and self.tri.area > 0 and self.trap.area > 0:
            print("\n" + "-"*50)
            print("ОПЕРАЦИИ С ФИГУРАМИ")
            print("-"*50)
            self.rect + self.tri
            self.tri - self.trap
            self.trap * 2
        
    def calc_trapezoid(self):
        self.trap.a = float(self.trap_a.get())
        self.trap.b = float(self.trap_b.get())
        self.trap.h = float(self.trap_h.get())
        result = self.trap.calculate()
        self.trap_result.config(text=f"Результат:\n{result}")
            
        str(self.trap)
        len(self.trap)
        repr(self.trap)
        if self.rect.area > 0 and self.tri.area > 0 and self.trap.area > 0:
            print("\n" + "-"*50)
            print("ОПЕРАЦИИ С ФИГУРАМИ")
            print("-"*50)
            self.rect + self.tri
            self.tri - self.trap
            self.trap * 2
  
    def save_to_word(self):
        doc = Document()
        doc.add_heading("Результаты расчётов", 0)
        
        if self.rect._result:
            doc.add_paragraph(f"Прямоугольник:\n{self.rect._result}\n")
        
        if self.tri._result:
            doc.add_paragraph(f"Треугольник:\n{self.tri._result}\n")
        
        if self.trap._result:
            doc.add_paragraph(f"Трапеция:\n{self.trap._result}\n")
        
        path = filedialog.asksaveasfilename(
            defaultextension=".docx", 
            initialfile="результаты.docx",
            filetypes=[("Word documents", "*.docx")]
        )
        if path:
            doc.save(path)
            messagebox.showinfo("Успех", "Файл сохранён!")
    
    def run(self):
        self.root.mainloop()

app = App()
app.run()